import re
import argparse
from docx import Document
import requests
import logging
from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException

# Logging configuration
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

file_handler = logging.FileHandler('translation_debug.log', 'w', 'utf-8')
console_handler = logging.StreamHandler()

file_handler.setLevel(logging.DEBUG)
console_handler.setLevel(logging.INFO)  # Output to console

formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

logger.addHandler(file_handler)
logger.addHandler(console_handler)

def detect_source_language(doc, min_words=50):
    """
    Detects the source language of the document by collecting text until reaching min_words.
    """
    collected_text = ''
    word_count = 0
    for para in doc.paragraphs:
        paragraph_text = ''.join(run.text for run in para.runs).strip()
        if paragraph_text:
            collected_text += ' ' + paragraph_text
            word_count += len(paragraph_text.split())
            if word_count >= min_words:
                break

    if word_count == 0:
        logger.error("Not enough text to detect source language.")
        return None

    try:
        src_lang = detect(collected_text)
        logger.info(f"Detected source language: {src_lang}")
        return src_lang
    except LangDetectException as e:
        logger.error(f"Language detection failed: {e}")
        return None

def call_ollama_api(text, src_lang, target_lang, model, api_token):
    api_url = 'http://localhost:11434/api/generate'

    prompt = f"""You are a professional translator like DeepL. Translate ONLY the text content from {src_lang} to {target_lang}.

Rules:
- Keep math formulas (e.g., LaTeX, equations), code, symbols, URLs, numbers, dates, proper names unchanged.
- Use precise technical terminology; preserve meaning, fluency, and structure (lists, emphasis).
- Output ONLY the translated text—no explanations, warnings, or extras.

Text: {text}"""

    payload = {
        'model': model,
        'prompt': prompt,
        'stream': False
    }

    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_token}'
    }

    try:
        logger.debug(f"Sending API request with payload: {payload}")
        response = requests.post(api_url, json=payload, headers=headers)
        logger.debug(f"Received response with status code {response.status_code}")
        if response.status_code == 200:
            response_json = response.json()
            logger.debug(f"API response JSON: {response_json}")
            translated_text = response_json.get('response', '').strip()
            logger.debug(f"Translated text: {translated_text}")
            return translated_text
        else:
            logger.error(f"API request failed with status code {response.status_code}")
            logger.error(f"Response: {response.text}")
            return text
    except Exception as e:
        logger.exception(f"An error occurred while calling the Ollama API: {e}")
        return text

def translate_text(text, src_lang, target_lang, model, api_token):
    logger.debug(f"Text to translate: '{text}'")

    if not text or not text.strip():
        logger.warning(f"Text is empty or whitespace: '{text}'")
        return text

    if src_lang == target_lang:
        logger.info(f"Source language and target language are the same for '{text}'. No translation needed.")
        return text

    translated_text = call_ollama_api(text, src_lang, target_lang, model, api_token)
    return translated_text

def process_paragraph(para, src_lang, model, target_lang, api_token):
    try:
        paragraph_text = ''.join(run.text for run in para.runs)
        logger.debug(f"Original paragraph text: '{paragraph_text}'")
        if paragraph_text.strip():
            translated_text = translate_text(paragraph_text, src_lang, target_lang, model, api_token)
            # Clear existing runs
            for run in para.runs:
                run.text = ''
            # Add a new run with the translated text
            para.add_run(translated_text)
            logger.debug(f"Translated paragraph text: '{translated_text}'")
        else:
            logger.debug("Paragraph is empty or whitespace.")
    except Exception as e:
        logger.exception(f"An error occurred while processing paragraph: {e}")

def process_document(input_file, output_file, model, target_lang, api_token, src_lang=None):
    try:
        logger.info(f"Opening document '{input_file}'")
        doc = Document(input_file)

        # Use provided source language or detect it
        if src_lang:
            logger.info(f"Using provided source language: {src_lang}")
        else:
            logger.info("Detecting source language...")
            src_lang = detect_source_language(doc)
            if not src_lang:
                logger.error("Source language detection failed. Exiting.")
                return

        logger.info("Processing main document paragraphs")
        for para in doc.paragraphs:
            process_paragraph(para, src_lang, model, target_lang, api_token)

        logger.info("Processing tables in the document")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para, src_lang, model, target_lang, api_token)

        logger.info("Processing headers and footers")
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for para in header.paragraphs:
                process_paragraph(para, src_lang, model, target_lang, api_token)
            for para in footer.paragraphs:
                process_paragraph(para, src_lang, model, target_lang, api_token)

        doc.save(output_file)
        logger.info(f"Document saved to '{output_file}'")
    except Exception as e:
        logger.exception(f"An error occurred while processing the document: {e}")

def main():
    parser = argparse.ArgumentParser(description='Translate a Word document using the Ollama API.')
    parser.add_argument('-m', '--model', type=str, default='llama3.2', help='The model name to use for translation.')
    parser.add_argument('-i', '--input_file', type=str, required=True, help='The input Word document file path.')
    parser.add_argument('-o', '--output_file', type=str, required=True, help='The output Word document file path.')
    parser.add_argument('-t', '--target_lang', type=str, required=True, help='The target language for translation (e.g., en, de, fr).')
    parser.add_argument('-k', '--api_token', type=str, required=True, help='Your API token for authentication.')
    parser.add_argument('-s', '--src_lang', type=str, help='The source language (e.g., en, de, fr). If not provided, the script will attempt to detect it.')
    args = parser.parse_args()

    logger.info("Starting translation process")
    process_document(args.input_file, args.output_file, args.model, args.target_lang, args.api_token, args.src_lang)
    logger.info("Translation process completed")

if __name__ == "__main__":
    main()
