from docx import Document

# Create a new document with intentional errors for proofreading test
doc = Document()

# Add paragraphs with grammar/spelling issues
doc.add_paragraph("This is a test docment for proofreading. It containts some grammer mistaks and spelling erros.")
doc.add_paragraph("The sentance structure could be beter in some places. Please fix these issues while keeping the meaning intact.")
doc.add_paragraph("Technical terms like API, URL, and LaTeX formulas (E = mc²) should remain unchanged.")

# Save it
doc.save('test_proofread.docx')
print("Created test_proofread.docx with intentional errors for proofreading testing")
