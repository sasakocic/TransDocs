"""
TransDocs - Document Translation and Proofreading Tool

A Python tool for translating Microsoft Word documents using the Ollama API.
Supports automatic language detection, professional translation quality,
and proofreading capabilities.
"""

import argparse
import os
from docx import Document
import requests
import logging
from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException


# ANSI color codes for colorful output
class ColorFormatter(logging.Formatter):
    """Custom formatter that adds colors to log levels."""

    COLORS = {
        "DEBUG": "\033[36m",  # Cyan
        "INFO": "\033[32m",  # Green
        "WARNING": "\033[33m",  # Yellow
        "ERROR": "\033[31m",  # Red
        "CRITICAL": "\033[35m",  # Magenta
    }

    RESET = "\033[0m"  # Reset color

    def format(self, record):
        log_color = self.COLORS.get(record.levelname, self.RESET)
        record.levelname = f"{log_color}{record.levelname}{self.RESET}"
        return super().format(record)


# Logging configuration
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

file_handler = logging.FileHandler("translation_debug.log", "w", encoding="utf-8")
console_handler = logging.StreamHandler()

# Use colorful formatter for console, plain text for file
console_formatter = ColorFormatter("%(asctime)s - %(levelname)s - %(message)s")
file_formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")

console_handler.setFormatter(console_formatter)
file_handler.setFormatter(file_formatter)

if not logger.handlers:
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)


def configure_logging(level_name="INFO"):
    """Configure logger and handlers with a shared log level."""
    level = getattr(logging, str(level_name).upper(), logging.INFO)
    logger.setLevel(level)
    file_handler.setLevel(level)
    console_handler.setLevel(level)


configure_logging(os.environ.get("TRANSDOC_LOG_LEVEL", "INFO"))


def detect_source_language(doc, min_words=50):
    """
    Detects the source language of the document by collecting text until reaching min_words.
    """
    collected_text = ""
    word_count = 0
    for para in doc.paragraphs:
        paragraph_text = "".join(run.text for run in para.runs).strip()
        if paragraph_text:
            collected_text += " " + paragraph_text
            word_count += len(paragraph_text.split())
            if word_count >= min_words:
                break

    if word_count == 0:
        logger.error("Not enough text to detect source language.")
        return None

    try:
        src_lang = detect(collected_text)
        logger.info(f"Detected source language: {src_lang}")
        return src_lang
    except LangDetectException as e:
        logger.error(f"Language detection failed: {e}")
        return None


def detect_language_from_blocks(text_blocks, min_words=50):
    """Detect language from extracted text blocks (e.g., PDF pages/paragraphs)."""
    collected_words = []
    for block in text_blocks:
        collected_words.extend(block.split())
        if len(collected_words) >= min_words:
            break

    if not collected_words:
        logger.error("Not enough text to detect source language.")
        return None

    try:
        sample = " ".join(collected_words)
        src_lang = detect(sample)
        logger.info(f"Detected source language: {src_lang}")
        return src_lang
    except LangDetectException as e:
        logger.error(f"Language detection failed: {e}")
        return None


def extract_pdf_text_blocks(input_file):
    """Extract text blocks from PDF for translation."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF support requires 'pypdf'. Install dependencies from requirements.txt."
        ) from exc

    reader = PdfReader(input_file)
    blocks = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        # Split large page text into paragraph-like blocks.
        page_blocks = [part.strip() for part in text.split("\n\n") if part.strip()]
        if page_blocks:
            blocks.extend(page_blocks)
        else:
            blocks.append(text)
    return blocks


def _load_pymupdf():
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise RuntimeError(
            "Layout-aware PDF translation requires 'PyMuPDF'. "
            "Install dependencies from requirements.txt."
        ) from exc
    return fitz


def extract_pdf_layout(input_file):
    """Extract page layout blocks for PDF translation with positional fidelity."""
    fitz = _load_pymupdf()
    source_pdf = fitz.open(input_file)
    pages = []
    plain_text_blocks = []

    try:
        for page in source_pdf:
            page_blocks = []
            for block in page.get_text("blocks"):
                x0, y0, x1, y1, text = block[:5]
                block_text = (text or "").strip()
                if not block_text:
                    continue
                page_blocks.append(
                    {
                        "bbox": (x0, y0, x1, y1),
                        "text": block_text,
                    }
                )
                plain_text_blocks.append(block_text)

            pages.append(
                {
                    "number": page.number,
                    "width": page.rect.width,
                    "height": page.rect.height,
                    "blocks": page_blocks,
                }
            )
    finally:
        source_pdf.close()

    return pages, plain_text_blocks


def _fit_textbox(page, rect, text, fitz):
    """Insert translated text in a rectangle with basic font-size fallback."""
    for size in (11, 10, 9, 8, 7):
        written = page.insert_textbox(
            rect,
            text,
            fontsize=size,
            fontname="helv",
            color=(0, 0, 0),
            align=fitz.TEXT_ALIGN_LEFT,
        )
        if written >= 0:
            return

    # Last resort: write clipped text near top-left of the block.
    page.insert_text(
        fitz.Point(rect.x0, rect.y0 + 8),
        text[:1000],
        fontsize=7,
        fontname="helv",
        color=(0, 0, 0),
    )


def translate_pdf_layout_to_pdf(
    input_file,
    output_file,
    model,
    target_lang,
    api_token,
    src_lang=None,
    api_url="http://localhost:11434",
    force_proofread=False,
    backend="ollama",
    progress_callback=None,
):
    """Translate PDF while preserving page geometry and text block placement."""
    fitz = _load_pymupdf()
    pages, text_blocks = extract_pdf_layout(input_file)
    if not text_blocks:
        raise ValueError("No extractable text found in PDF")

    if src_lang:
        logger.info(f"Using provided source language: {src_lang}")
    else:
        logger.info("Detecting source language...")
        src_lang = detect_language_from_blocks(text_blocks)
        if not src_lang:
            logger.error("Source language detection failed. Exiting.")
            return

    if force_proofread:
        logger.info("Running in PROOFREADING mode (explicit --proofread flag).")
    elif src_lang == target_lang:
        logger.info(
            f"Source ({src_lang}) equals target ({target_lang}). Running in PROOFREADING mode."
        )
    else:
        logger.info(f"Translating from {src_lang} to {target_lang}.")

    total_units = sum(len(page_data["blocks"]) for page_data in pages)
    logger.info(f"Processing {total_units} positioned text blocks from PDF")
    if progress_callback:
        progress_callback(0, total_units, "Starting translation")

    source_pdf = fitz.open(input_file)
    output_pdf = fitz.open()

    try:
        completed = 0
        for page_data in pages:
            output_page = output_pdf.new_page(
                width=page_data["width"], height=page_data["height"]
            )
            # Keep original page graphics/layout as background.
            output_page.show_pdf_page(output_page.rect, source_pdf, page_data["number"])

            translated_blocks = []
            for block in page_data["blocks"]:
                translated_text = translate_or_proofread(
                    block["text"],
                    src_lang,
                    target_lang,
                    model,
                    api_token,
                    api_url,
                    force_proofread=force_proofread,
                    backend=backend,
                )
                translated_blocks.append((fitz.Rect(block["bbox"]), translated_text))
                completed += 1
                if progress_callback:
                    progress_callback(
                        completed, total_units, f"Processing {completed}/{total_units}"
                    )

            # Redact original text blocks, then insert translated text at same positions.
            for rect, _ in translated_blocks:
                output_page.add_redact_annot(rect, fill=(1, 1, 1))
            output_page.apply_redactions()

            for rect, translated_text in translated_blocks:
                _fit_textbox(output_page, rect, translated_text, fitz)

        output_pdf.save(output_file, garbage=4, deflate=True)
        logger.info(f"Document saved to '{output_file}'")
        if progress_callback:
            progress_callback(total_units, total_units, "Completed")
    finally:
        source_pdf.close()
        output_pdf.close()


def build_chat_endpoints(api_url, backend):
    """Build ordered chat endpoint candidates for the selected backend."""
    base_url = api_url.rstrip("/")
    endpoints = []

    def add(url):
        if url not in endpoints:
            endpoints.append(url)

    if backend == "openai_compatible":
        if base_url.endswith(("/v1/chat/completions", "/chat/completions")):
            add(base_url)
        elif base_url.endswith("/v1"):
            add(f"{base_url}/chat/completions")
        else:
            add(f"{base_url}/v1/chat/completions")
        # Fallback for gateways exposing Ollama-style endpoints
        if base_url.endswith("/api/chat"):
            add(base_url)
        else:
            add(f"{base_url}/api/chat")
        return endpoints

    # Default: ollama
    if base_url.endswith(("/api/chat", "/api/generate")):
        add(base_url)
    else:
        add(f"{base_url}/api/chat")
    # Fallback for gateways exposing OpenAI-compatible endpoints
    if base_url.endswith(("/v1/chat/completions", "/chat/completions")):
        add(base_url)
    elif base_url.endswith("/v1"):
        add(f"{base_url}/chat/completions")
    else:
        add(f"{base_url}/v1/chat/completions")
    return endpoints


def build_chat_endpoint(api_url, backend):
    """Backward-compatible single endpoint helper (first candidate)."""
    return build_chat_endpoints(api_url, backend)[0]


def _extract_response_text(response_json):
    """Extract assistant text from Ollama or OpenAI-compatible responses."""
    if "response" in response_json:
        return (response_json.get("response") or "").strip()

    message = response_json.get("message", {})
    if isinstance(message, dict) and "content" in message:
        return (message.get("content") or "").strip()

    choices = response_json.get("choices", [])
    if isinstance(choices, list) and choices:
        first_choice = choices[0] or {}
        choice_message = first_choice.get("message", {})
        if isinstance(choice_message, dict) and "content" in choice_message:
            return (choice_message.get("content") or "").strip()
        if "text" in first_choice:
            return (first_choice.get("text") or "").strip()

    return ""


def call_chat_api(
    text,
    src_lang,
    target_lang,
    model,
    api_token,
    api_url,
    mode="translate",
    backend="ollama",
):
    if mode == "proofread":
        prompt = f"""You are a professional proofreader and editor. Review the following text in {src_lang} for grammar, spelling, punctuation, clarity, and style improvements.

Rules:
- Keep math formulas (e.g., LaTeX, equations), code, symbols, URLs, numbers, dates, proper names unchanged.
- Preserve the original meaning while improving fluency, readability, and correctness.
- Fix grammatical errors, typos, and awkward phrasing.
- Maintain formatting, lists, emphasis, and document structure.
- Preserve line breaks exactly when possible; do not collapse multi-line structure into a single paragraph.
- Keep output length close to source length to help preserve document/PDF layout.
- Do not add headings, explanations, or extra sentences.
- Output ONLY the improved text—no explanations, comments, or notes about changes made.

Text: {text}"""
    else:  # translate mode
        prompt = f"""You are a professional translator like DeepL. Translate ONLY the text content from {src_lang} to {target_lang}.

Rules:
- Keep math formulas (e.g., LaTeX, equations), code, symbols, URLs, numbers, dates, proper names unchanged.
- Use precise technical terminology; preserve meaning, fluency, and structure (lists, emphasis).
- Preserve line breaks and list structure exactly when possible.
- Keep translated text compact and similar in length to source text to improve layout fidelity.
- Avoid adding or removing sentences unless strictly required by grammar.
- Output ONLY the translated text—no explanations, warnings, or extras.

Text: {text}"""

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "stream": False,
    }

    headers = {"Content-Type": "application/json"}

    # Add authorization header only if token is provided
    if api_token:
        headers["Authorization"] = f"Bearer {api_token}"
        headers["X-API-Key"] = api_token
        headers["api-key"] = api_token

    try:
        endpoint_candidates = build_chat_endpoints(api_url, backend)
        last_error = None
        last_status = None
        last_response_text = ""

        for endpoint_url in endpoint_candidates:
            # Extract just the endpoint path for logging
            import urllib.parse

            parsed_url = urllib.parse.urlparse(endpoint_url)
            endpoint_path = parsed_url.path if parsed_url.path else "/"

            logger.debug("=== CHAT API CALL DEBUG ===")
            logger.debug(f"Backend: {backend}")
            logger.debug(f"Full URL: {endpoint_url}")
            logger.debug(f"Endpoint path: {endpoint_path}")
            logger.debug(f"Model: {model}")
            logger.debug(f"Messages count: {len(payload.get('messages', []))}")
            logger.debug(
                f"Prompt preview (first 200 chars): {payload.get('messages', [{}])[0].get('content', '')[:200]}"
            )
            logger.debug(f"Request headers: {headers}")
            logger.debug("HTTP method being used: POST")

            response = requests.post(endpoint_url, json=payload, headers=headers)
            logger.debug(f"Response status code: {response.status_code}")
            logger.debug(f"Response headers: {dict(response.headers)}")
            logger.debug(f"Response body (raw): {response.text[:500]}")

            # Log full response if error to help debug
            if response.status_code != 200:
                logger.error(f"Full response text: {response.text}")

            if response.status_code == 200:
                response_json = response.json()
                logger.debug(f"API response JSON: {response_json}")

                result_text = _extract_response_text(response_json)
                logger.debug(f"{mode.capitalize()} text: {result_text}")
                return result_text.strip() if result_text else ""

            last_status = response.status_code
            last_response_text = response.text
            last_error = f"HTTP {response.status_code}"

        logger.error(
            f"API request failed after trying {len(endpoint_candidates)} endpoints"
        )
        if last_status:
            logger.error(f"Last status: {last_status}")
        if last_response_text:
            logger.error(f"Last response: {last_response_text}")
        if last_error:
            logger.error(f"Error: {last_error}")
        return text
    except Exception as e:
        logger.exception(f"An error occurred while calling the chat API: {e}")
        return text


def call_ollama_api(
    text, src_lang, target_lang, model, api_token, api_url, mode="translate"
):
    """Backward-compatible wrapper for Ollama-only calls."""
    return call_chat_api(
        text,
        src_lang,
        target_lang,
        model,
        api_token,
        api_url,
        mode=mode,
        backend="ollama",
    )


def translate_or_proofread(
    text,
    src_lang,
    target_lang,
    model,
    api_token,
    api_url,
    force_proofread=False,
    backend="ollama",
):
    logger.debug(f"Text to process: '{text}'")

    if not text or not text.strip():
        logger.warning("Text is empty or whitespace.")
        return text

    # Determine mode based on explicit flag or language match
    if force_proofread:
        mode = "proofread"
        logger.info("Forced proofreading mode (explicit --proofread flag).")
    elif src_lang == target_lang:
        mode = "proofread"
        logger.info(
            f"Source and target languages are the same ({src_lang}). Running in proofreading mode."
        )
    else:
        mode = "translate"
        logger.info(f"Translating from {src_lang} to {target_lang}.")

    result_text = call_chat_api(
        text,
        src_lang,
        target_lang,
        model,
        api_token,
        api_url,
        mode=mode,
        backend=backend,
    )
    return result_text


def process_paragraph(
    para,
    src_lang,
    model,
    target_lang,
    api_token,
    api_url,
    force_proofread=False,
    backend="ollama",
):
    try:
        paragraph_text = "".join(run.text for run in para.runs)
        logger.debug(f"Original paragraph text: '{paragraph_text}'")
        if paragraph_text.strip():
            result_text = translate_or_proofread(
                paragraph_text,
                src_lang,
                target_lang,
                model,
                api_token,
                api_url,
                force_proofread=force_proofread,
                backend=backend,
            )
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            # Add a new run with the processed text
            para.add_run(result_text)
            logger.debug(f"Processed paragraph text: '{result_text}'")
        else:
            logger.debug("Paragraph is empty or whitespace.")
    except Exception as e:
        logger.exception(f"An error occurred while processing paragraph: {e}")


def iter_document_paragraphs(doc):
    """Yield all paragraphs that are processed (body, tables, headers, footers)."""
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for para in section.footer.paragraphs:
            yield para


def process_document(
    input_file,
    output_file,
    model,
    target_lang,
    api_token,
    src_lang=None,
    api_url="http://localhost:11434",
    force_proofread=False,
    backend="ollama",
    progress_callback=None,
):
    try:
        logger.info(f"Opening document '{input_file}'")
        extension = os.path.splitext(input_file)[1].lower()
        output_extension = os.path.splitext(output_file)[1].lower()

        if extension == ".pdf":
            if output_extension == ".pdf":
                logger.info(
                    "Detected PDF input/output. Using layout-aware PDF translation."
                )
                translate_pdf_layout_to_pdf(
                    input_file,
                    output_file,
                    model,
                    target_lang,
                    api_token,
                    src_lang=src_lang,
                    api_url=api_url,
                    force_proofread=force_proofread,
                    backend=backend,
                    progress_callback=progress_callback,
                )
                return

            logger.info("Detected PDF input. Extracting text blocks...")
            text_blocks = extract_pdf_text_blocks(input_file)
            if not text_blocks:
                raise ValueError("No extractable text found in PDF")

            if src_lang:
                logger.info(f"Using provided source language: {src_lang}")
            else:
                logger.info("Detecting source language...")
                src_lang = detect_language_from_blocks(text_blocks)
                if not src_lang:
                    logger.error("Source language detection failed. Exiting.")
                    return

            if force_proofread:
                logger.info("Running in PROOFREADING mode (explicit --proofread flag).")
            elif src_lang == target_lang:
                logger.info(
                    f"Source ({src_lang}) equals target ({target_lang}). Running in PROOFREADING mode."
                )
            else:
                logger.info(f"Translating from {src_lang} to {target_lang}.")

            total_units = len(text_blocks)
            if progress_callback:
                progress_callback(0, total_units, "Starting translation")

            output_doc = Document()
            logger.info(f"Processing {total_units} text blocks from PDF")
            for index, block in enumerate(text_blocks, start=1):
                result_text = translate_or_proofread(
                    block,
                    src_lang,
                    target_lang,
                    model,
                    api_token,
                    api_url,
                    force_proofread=force_proofread,
                    backend=backend,
                )
                output_doc.add_paragraph(result_text)
                if progress_callback:
                    progress_callback(
                        index, total_units, f"Processing {index}/{total_units}"
                    )

            output_doc.save(output_file)
            logger.info(f"Document saved to '{output_file}'")
            if progress_callback:
                progress_callback(total_units, total_units, "Completed")
            return

        if extension != ".docx":
            raise ValueError("Unsupported input format. Use .docx or .pdf")

        doc = Document(input_file)

        # Use provided source language or detect it
        if src_lang:
            logger.info(f"Using provided source language: {src_lang}")
        else:
            logger.info("Detecting source language...")
            src_lang = detect_source_language(doc)
            if not src_lang:
                logger.error("Source language detection failed. Exiting.")
                return

        # Check if we need proofreading or translation
        if force_proofread:
            logger.info("Running in PROOFREADING mode (explicit --proofread flag).")
        elif src_lang == target_lang:
            logger.info(
                f"Source ({src_lang}) equals target ({target_lang}). Running in PROOFREADING mode."
            )
        else:
            logger.info(f"Translating from {src_lang} to {target_lang}.")

        all_paragraphs = list(iter_document_paragraphs(doc))
        total_units = len(all_paragraphs)
        logger.info(f"Processing {total_units} text blocks")
        if progress_callback:
            progress_callback(0, total_units, "Starting translation")

        for index, para in enumerate(all_paragraphs, start=1):
            process_paragraph(
                para,
                src_lang,
                model,
                target_lang,
                api_token,
                api_url,
                force_proofread=force_proofread,
                backend=backend,
            )
            if progress_callback:
                progress_callback(
                    index, total_units, f"Processing {index}/{total_units}"
                )

        doc.save(output_file)
        logger.info(f"Document saved to '{output_file}'")
        if progress_callback:
            progress_callback(total_units, total_units, "Completed")
    except Exception as e:
        logger.exception(f"An error occurred while processing the document: {e}")
        raise


def main():
    parser = argparse.ArgumentParser(
        description="Translate or proofread a Word document using Ollama or OpenAI-compatible chat APIs."
    )
    parser.add_argument(
        "-m",
        "--model",
        type=str,
        default="",
        help="Model name (required). Leave empty to error out if not provided.",
    )
    parser.add_argument(
        "-i",
        "--input",
        type=str,
        required=True,
        dest="input_file",
        metavar="FILE",
        help="Path to the input Word document file.",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=str,
        required=True,
        dest="output_file",
        metavar="FILE",
        help="Path to save the output Word document file.",
    )
    parser.add_argument(
        "-t",
        "--target",
        type=str,
        required=True,
        dest="target_lang",
        metavar="LANG",
        help="Target language code (e.g., en, de, fr). Use same as source for proofreading mode.",
    )
    parser.add_argument(
        "-k",
        "--api-token",
        type=str,
        default=None,
        dest="api_token",
        metavar="TOKEN",
        help="API token for Ollama authentication (optional).",
    )
    parser.add_argument(
        "-s",
        "--source",
        type=str,
        default=None,
        dest="src_lang",
        metavar="LANG",
        help="Source language code (e.g., en, de, fr). Auto-detected if not provided. Use same as target for proofreading mode.",
    )
    parser.add_argument(
        "--proofread",
        action="store_true",
        help="Force proofreading mode regardless of language match.",
    )
    parser.add_argument(
        "-u",
        "--url",
        type=str,
        default="http://localhost:11434",
        dest="api_url",
        metavar="URL",
        help="Ollama API base URL (default: http://localhost:11434). Auto-appends /api/chat.",
    )
    parser.add_argument(
        "-b",
        "--backend",
        choices=["ollama", "openai_compatible"],
        default="ollama",
        help="API backend type (default: ollama).",
    )
    parser.add_argument(
        "-v",
        "--verbose",
        action="count",
        default=0,
        help="Enable verbose logging. Use -v for DEBUG logs.",
    )
    parser.add_argument(
        "--log-level",
        type=str,
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        default=None,
        help="Explicit log level override (default: INFO).",
    )
    args = parser.parse_args()

    if args.log_level:
        log_level = args.log_level
    elif args.verbose >= 1:
        log_level = "DEBUG"
    else:
        log_level = os.environ.get("TRANSDOC_LOG_LEVEL", "INFO")
    configure_logging(log_level)

    # Validate model is provided
    if not args.model or not args.model.strip():
        parser.error("Model name (-m/--model) is required")

    api_url = build_chat_endpoint(args.api_url, args.backend)

    logger.info(
        f"Starting processing with backend '{args.backend}' and API URL: {api_url}"
    )
    process_document(
        args.input_file,
        args.output_file,
        args.model,
        args.target_lang,
        args.api_token,
        args.src_lang,
        api_url,
        force_proofread=args.proofread,
        backend=args.backend,
    )
    logger.info("Processing completed")


if __name__ == "__main__":
    main()
